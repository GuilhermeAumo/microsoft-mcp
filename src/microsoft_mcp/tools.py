import base64
import datetime as dt
import io
import pathlib as pl
from typing import Any
from docx import Document
from fastmcp import FastMCP
from . import graph, auth

mcp = FastMCP("microsoft-mcp")

FOLDERS = {
    k.casefold(): v
    for k, v in {
        "inbox": "inbox",
        "sent": "sentitems",
        "drafts": "drafts",
        "deleted": "deleteditems",
        "junk": "junkemail",
        "archive": "archive",
    }.items()
}


@mcp.tool
def list_accounts() -> list[dict[str, str]]:
    """List all signed-in Microsoft accounts"""
    return [
        {"username": acc.username, "account_id": acc.account_id}
        for acc in auth.list_accounts()
    ]


@mcp.tool
def authenticate_account() -> dict[str, str]:
    """Authenticate a new Microsoft account using device flow authentication

    Returns authentication instructions and device code for the user to complete authentication.
    The user must visit the URL and enter the code to authenticate their Microsoft account.
    """
    app = auth.get_app()
    flow = app.initiate_device_flow(scopes=auth.SCOPES)

    if "user_code" not in flow:
        error_msg = flow.get("error_description", "Unknown error")
        raise Exception(f"Failed to get device code: {error_msg}")

    verification_url = flow.get(
        "verification_uri",
        flow.get("verification_url", "https://microsoft.com/devicelogin"),
    )

    return {
        "status": "authentication_required",
        "instructions": "To authenticate a new Microsoft account:",
        "step1": f"Visit: {verification_url}",
        "step2": f"Enter code: {flow['user_code']}",
        "step3": "Sign in with the Microsoft account you want to add",
        "step4": "After authenticating, use the 'complete_authentication' tool to finish the process",
        "device_code": flow["user_code"],
        "verification_url": verification_url,
        "expires_in": flow.get("expires_in", 900),
        "_flow_cache": str(flow),
    }


@mcp.tool
def complete_authentication(flow_cache: str) -> dict[str, str]:
    """Complete the authentication process after the user has entered the device code

    Args:
        flow_cache: The flow data returned from authenticate_account (the _flow_cache field)

    Returns:
        Account information if authentication was successful
    """
    import ast

    try:
        flow = ast.literal_eval(flow_cache)
    except (ValueError, SyntaxError):
        raise ValueError("Invalid flow cache data")

    app = auth.get_app()
    result = app.acquire_token_by_device_flow(flow)

    if "error" in result:
        error_msg = result.get("error_description", result["error"])
        if "authorization_pending" in error_msg:
            return {
                "status": "pending",
                "message": "Authentication is still pending. The user needs to complete the authentication process.",
                "instructions": "Please ensure you've visited the URL and entered the code, then try again.",
            }
        raise Exception(f"Authentication failed: {error_msg}")

    # Save the token cache
    cache = app.token_cache
    if isinstance(cache, auth.msal.SerializableTokenCache) and cache.has_state_changed:
        auth._write_cache(cache.serialize())

    # Get the newly added account
    accounts = app.get_accounts()
    if accounts:
        # Find the account that matches the token we just got
        for account in accounts:
            if (
                account.get("username", "").lower()
                == result.get("id_token_claims", {})
                .get("preferred_username", "")
                .lower()
            ):
                return {
                    "status": "success",
                    "username": account["username"],
                    "account_id": account["home_account_id"],
                    "message": f"Successfully authenticated {account['username']}",
                }
        # If exact match not found, return the last account
        account = accounts[-1]
        return {
            "status": "success",
            "username": account["username"],
            "account_id": account["home_account_id"],
            "message": f"Successfully authenticated {account['username']}",
        }

    return {
        "status": "error",
        "message": "Authentication succeeded but no account was found",
    }


@mcp.tool
def list_emails(
    account_id: str,
    folder: str = "inbox",
    limit: int = 10,
    include_body: bool = True,
) -> list[dict[str, Any]]:
    """List emails from specified folder"""
    folder_path = FOLDERS.get(folder.casefold(), folder)

    if include_body:
        select_fields = "id,subject,from,toRecipients,ccRecipients,receivedDateTime,hasAttachments,body,conversationId,isRead"
    else:
        select_fields = "id,subject,from,toRecipients,receivedDateTime,hasAttachments,conversationId,isRead"

    params = {
        "$top": min(limit, 100),
        "$select": select_fields,
        "$orderby": "receivedDateTime desc",
    }

    emails = list(
        graph.request_paginated(
            f"/me/mailFolders/{folder_path}/messages",
            account_id,
            params=params,
            limit=limit,
        )
    )

    return emails


@mcp.tool
def get_email(
    email_id: str,
    account_id: str,
    include_body: bool = True,
    body_max_length: int = 50000,
    include_attachments: bool = True,
) -> dict[str, Any]:
    """Get email details with size limits

    Args:
        email_id: The email ID
        account_id: The account ID
        include_body: Whether to include the email body (default: True)
        body_max_length: Maximum characters for body content (default: 50000)
        include_attachments: Whether to include attachment metadata (default: True)
    """
    params = {}
    if include_attachments:
        params["$expand"] = "attachments($select=id,name,size,contentType)"

    result = graph.request("GET", f"/me/messages/{email_id}", account_id, params=params)
    if not result:
        raise ValueError(f"Email with ID {email_id} not found")

    # Truncate body if needed
    if include_body and "body" in result and "content" in result["body"]:
        content = result["body"]["content"]
        if len(content) > body_max_length:
            result["body"]["content"] = (
                content[:body_max_length]
                + f"\n\n[Content truncated - {len(content)} total characters]"
            )
            result["body"]["truncated"] = True
            result["body"]["total_length"] = len(content)
    elif not include_body and "body" in result:
        del result["body"]

    # Remove attachment content bytes to reduce size
    if "attachments" in result and result["attachments"]:
        for attachment in result["attachments"]:
            if "contentBytes" in attachment:
                del attachment["contentBytes"]

    return result


@mcp.tool
def create_email_draft(
    account_id: str,
    to: str | list[str],
    subject: str,
    body: str,
    cc: str | list[str] | None = None,
    attachments: str | list[str] | None = None,
) -> dict[str, Any]:
    """Create an email draft with file path(s) as attachments"""
    to_list = [to] if isinstance(to, str) else to

    message = {
        "subject": subject,
        "body": {"contentType": "Text", "content": body},
        "toRecipients": [{"emailAddress": {"address": addr}} for addr in to_list],
    }

    if cc:
        cc_list = [cc] if isinstance(cc, str) else cc
        message["ccRecipients"] = [
            {"emailAddress": {"address": addr}} for addr in cc_list
        ]

    small_attachments = []
    large_attachments = []

    if attachments:
        # Convert single path to list
        attachment_paths = (
            [attachments] if isinstance(attachments, str) else attachments
        )
        for file_path in attachment_paths:
            path = pl.Path(file_path).expanduser().resolve()
            content_bytes = path.read_bytes()
            att_size = len(content_bytes)
            att_name = path.name

            if att_size < 3 * 1024 * 1024:
                small_attachments.append(
                    {
                        "@odata.type": "#microsoft.graph.fileAttachment",
                        "name": att_name,
                        "contentBytes": base64.b64encode(content_bytes).decode("utf-8"),
                    }
                )
            else:
                large_attachments.append(
                    {
                        "name": att_name,
                        "content_bytes": content_bytes,
                        "content_type": "application/octet-stream",
                    }
                )

    if small_attachments:
        message["attachments"] = small_attachments

    result = graph.request("POST", "/me/messages", account_id, json=message)
    if not result:
        raise ValueError("Failed to create email draft")

    message_id = result["id"]

    for att in large_attachments:
        graph.upload_large_mail_attachment(
            message_id,
            att["name"],
            att["content_bytes"],
            account_id,
            att.get("content_type", "application/octet-stream"),
        )

    return result


@mcp.tool
def send_email(
    account_id: str,
    to: str | list[str],
    subject: str,
    body: str,
    cc: str | list[str] | None = None,
    attachments: str | list[str] | None = None,
) -> dict[str, str]:
    """Send an email immediately with file path(s) as attachments"""
    to_list = [to] if isinstance(to, str) else to

    message = {
        "subject": subject,
        "body": {"contentType": "Text", "content": body},
        "toRecipients": [{"emailAddress": {"address": addr}} for addr in to_list],
    }

    if cc:
        cc_list = [cc] if isinstance(cc, str) else cc
        message["ccRecipients"] = [
            {"emailAddress": {"address": addr}} for addr in cc_list
        ]

    # Check if we have large attachments
    has_large_attachments = False
    processed_attachments = []

    if attachments:
        # Convert single path to list
        attachment_paths = (
            [attachments] if isinstance(attachments, str) else attachments
        )
        for file_path in attachment_paths:
            path = pl.Path(file_path).expanduser().resolve()
            content_bytes = path.read_bytes()
            att_size = len(content_bytes)
            att_name = path.name

            processed_attachments.append(
                {
                    "name": att_name,
                    "content_bytes": content_bytes,
                    "content_type": "application/octet-stream",
                    "size": att_size,
                }
            )

            if att_size >= 3 * 1024 * 1024:
                has_large_attachments = True

    if not has_large_attachments and processed_attachments:
        message["attachments"] = [
            {
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": att["name"],
                "contentBytes": base64.b64encode(att["content_bytes"]).decode("utf-8"),
            }
            for att in processed_attachments
        ]
        graph.request("POST", "/me/sendMail", account_id, json={"message": message})
        return {"status": "sent"}
    elif has_large_attachments:
        # Create draft first, then add large attachments, then send
        # We need to handle large attachments manually here
        to_list = [to] if isinstance(to, str) else to
        message = {
            "subject": subject,
            "body": {"contentType": "Text", "content": body},
            "toRecipients": [{"emailAddress": {"address": addr}} for addr in to_list],
        }
        if cc:
            cc_list = [cc] if isinstance(cc, str) else cc
            message["ccRecipients"] = [
                {"emailAddress": {"address": addr}} for addr in cc_list
            ]

        result = graph.request("POST", "/me/messages", account_id, json=message)
        if not result:
            raise ValueError("Failed to create email draft")

        message_id = result["id"]

        for att in processed_attachments:
            if att["size"] >= 3 * 1024 * 1024:
                graph.upload_large_mail_attachment(
                    message_id,
                    att["name"],
                    att["content_bytes"],
                    account_id,
                    att.get("content_type", "application/octet-stream"),
                )
            else:
                small_att = {
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": att["name"],
                    "contentBytes": base64.b64encode(att["content_bytes"]).decode(
                        "utf-8"
                    ),
                }
                graph.request(
                    "POST",
                    f"/me/messages/{message_id}/attachments",
                    account_id,
                    json=small_att,
                )

        graph.request("POST", f"/me/messages/{message_id}/send", account_id)
        return {"status": "sent"}
    else:
        graph.request("POST", "/me/sendMail", account_id, json={"message": message})
        return {"status": "sent"}


@mcp.tool
def update_email(
    email_id: str, updates: dict[str, Any], account_id: str
) -> dict[str, Any]:
    """Update email properties (isRead, categories, flag, etc.)"""
    result = graph.request(
        "PATCH", f"/me/messages/{email_id}", account_id, json=updates
    )
    if not result:
        raise ValueError(f"Failed to update email {email_id} - no response")
    return result


@mcp.tool
def delete_email(email_id: str, account_id: str) -> dict[str, str]:
    """Delete an email"""
    graph.request("DELETE", f"/me/messages/{email_id}", account_id)
    return {"status": "deleted"}


@mcp.tool
def move_email(
    email_id: str, destination_folder: str, account_id: str
) -> dict[str, Any]:
    """Move email to another folder"""
    folder_path = FOLDERS.get(destination_folder.casefold(), destination_folder)

    folders = graph.request("GET", "/me/mailFolders", account_id)
    folder_id = None

    if not folders:
        raise ValueError("Failed to retrieve mail folders")
    if "value" not in folders:
        raise ValueError(f"Unexpected folder response structure: {folders}")

    for folder in folders["value"]:
        if folder["displayName"].lower() == folder_path.lower():
            folder_id = folder["id"]
            break

    if not folder_id:
        raise ValueError(f"Folder '{destination_folder}' not found")

    payload = {"destinationId": folder_id}
    result = graph.request(
        "POST", f"/me/messages/{email_id}/move", account_id, json=payload
    )
    if not result:
        raise ValueError("Failed to move email - no response from server")
    if "id" not in result:
        raise ValueError(f"Failed to move email - unexpected response: {result}")
    return {"status": "moved", "new_id": result["id"]}


@mcp.tool
def reply_to_email(account_id: str, email_id: str, body: str) -> dict[str, str]:
    """Reply to an email (sender only)"""
    endpoint = f"/me/messages/{email_id}/reply"
    payload = {"message": {"body": {"contentType": "Text", "content": body}}}
    graph.request("POST", endpoint, account_id, json=payload)
    return {"status": "sent"}


@mcp.tool
def reply_all_email(account_id: str, email_id: str, body: str) -> dict[str, str]:
    """Reply to all recipients of an email"""
    endpoint = f"/me/messages/{email_id}/replyAll"
    payload = {"message": {"body": {"contentType": "Text", "content": body}}}
    graph.request("POST", endpoint, account_id, json=payload)
    return {"status": "sent"}


@mcp.tool
def list_events(
    account_id: str,
    days_ahead: int = 7,
    days_back: int = 0,
    include_details: bool = True,
) -> list[dict[str, Any]]:
    """List calendar events within specified date range, including recurring event instances"""
    now = dt.datetime.now(dt.timezone.utc)
    start = (now - dt.timedelta(days=days_back)).isoformat()
    end = (now + dt.timedelta(days=days_ahead)).isoformat()

    params = {
        "startDateTime": start,
        "endDateTime": end,
        "$orderby": "start/dateTime",
        "$top": 100,
    }

    if include_details:
        params["$select"] = (
            "id,subject,start,end,location,body,attendees,organizer,isAllDay,recurrence,onlineMeeting,seriesMasterId"
        )
    else:
        params["$select"] = "id,subject,start,end,location,organizer,seriesMasterId"

    # Use calendarView to get recurring event instances
    events = list(
        graph.request_paginated("/me/calendarView", account_id, params=params)
    )

    return events


@mcp.tool
def get_event(event_id: str, account_id: str) -> dict[str, Any]:
    """Get full event details"""
    result = graph.request("GET", f"/me/events/{event_id}", account_id)
    if not result:
        raise ValueError(f"Event with ID {event_id} not found")
    return result


@mcp.tool
def create_event(
    account_id: str,
    subject: str,
    start: str,
    end: str,
    location: str | None = None,
    body: str | None = None,
    attendees: str | list[str] | None = None,
    timezone: str = "UTC",
) -> dict[str, Any]:
    """Create a calendar event"""
    event = {
        "subject": subject,
        "start": {"dateTime": start, "timeZone": timezone},
        "end": {"dateTime": end, "timeZone": timezone},
    }

    if location:
        event["location"] = {"displayName": location}

    if body:
        event["body"] = {"contentType": "Text", "content": body}

    if attendees:
        attendees_list = [attendees] if isinstance(attendees, str) else attendees
        event["attendees"] = [
            {"emailAddress": {"address": a}, "type": "required"} for a in attendees_list
        ]

    result = graph.request("POST", "/me/events", account_id, json=event)
    if not result:
        raise ValueError("Failed to create event")
    return result


@mcp.tool
def update_event(
    event_id: str, updates: dict[str, Any], account_id: str
) -> dict[str, Any]:
    """Update event properties"""
    formatted_updates = {}

    if "subject" in updates:
        formatted_updates["subject"] = updates["subject"]
    if "start" in updates:
        formatted_updates["start"] = {
            "dateTime": updates["start"],
            "timeZone": updates.get("timezone", "UTC"),
        }
    if "end" in updates:
        formatted_updates["end"] = {
            "dateTime": updates["end"],
            "timeZone": updates.get("timezone", "UTC"),
        }
    if "location" in updates:
        formatted_updates["location"] = {"displayName": updates["location"]}
    if "body" in updates:
        formatted_updates["body"] = {"contentType": "Text", "content": updates["body"]}

    result = graph.request(
        "PATCH", f"/me/events/{event_id}", account_id, json=formatted_updates
    )
    return result or {"status": "updated"}


@mcp.tool
def delete_event(
    account_id: str, event_id: str, send_cancellation: bool = True
) -> dict[str, str]:
    """Delete or cancel a calendar event"""
    if send_cancellation:
        graph.request("POST", f"/me/events/{event_id}/cancel", account_id, json={})
    else:
        graph.request("DELETE", f"/me/events/{event_id}", account_id)
    return {"status": "deleted"}


@mcp.tool
def respond_event(
    account_id: str,
    event_id: str,
    response: str = "accept",
    message: str | None = None,
) -> dict[str, str]:
    """Respond to event invitation (accept, decline, tentativelyAccept)"""
    payload: dict[str, Any] = {"sendResponse": True}
    if message:
        payload["comment"] = message

    graph.request("POST", f"/me/events/{event_id}/{response}", account_id, json=payload)
    return {"status": response}


@mcp.tool
def check_availability(
    account_id: str,
    start: str,
    end: str,
    attendees: str | list[str] | None = None,
) -> dict[str, Any]:
    """Check calendar availability for scheduling"""
    me_info = graph.request("GET", "/me", account_id)
    if not me_info or "mail" not in me_info:
        raise ValueError("Failed to get user email address")
    schedules = [me_info["mail"]]
    if attendees:
        attendees_list = [attendees] if isinstance(attendees, str) else attendees
        schedules.extend(attendees_list)

    payload = {
        "schedules": schedules,
        "startTime": {"dateTime": start, "timeZone": "UTC"},
        "endTime": {"dateTime": end, "timeZone": "UTC"},
        "availabilityViewInterval": 30,
    }

    result = graph.request("POST", "/me/calendar/getSchedule", account_id, json=payload)
    if not result:
        raise ValueError("Failed to check availability")
    return result


@mcp.tool
def list_contacts(account_id: str, limit: int = 50) -> list[dict[str, Any]]:
    """List contacts"""
    params = {"$top": min(limit, 100)}

    contacts = list(
        graph.request_paginated("/me/contacts", account_id, params=params, limit=limit)
    )

    return contacts


@mcp.tool
def get_contact(contact_id: str, account_id: str) -> dict[str, Any]:
    """Get contact details"""
    result = graph.request("GET", f"/me/contacts/{contact_id}", account_id)
    if not result:
        raise ValueError(f"Contact with ID {contact_id} not found")
    return result


@mcp.tool
def create_contact(
    account_id: str,
    given_name: str,
    surname: str | None = None,
    email_addresses: str | list[str] | None = None,
    phone_numbers: dict[str, str] | None = None,
) -> dict[str, Any]:
    """Create a new contact"""
    contact: dict[str, Any] = {"givenName": given_name}

    if surname:
        contact["surname"] = surname

    if email_addresses:
        email_list = (
            [email_addresses] if isinstance(email_addresses, str) else email_addresses
        )
        contact["emailAddresses"] = [
            {"address": email, "name": f"{given_name} {surname or ''}".strip()}
            for email in email_list
        ]

    if phone_numbers:
        if "business" in phone_numbers:
            contact["businessPhones"] = [phone_numbers["business"]]
        if "home" in phone_numbers:
            contact["homePhones"] = [phone_numbers["home"]]
        if "mobile" in phone_numbers:
            contact["mobilePhone"] = phone_numbers["mobile"]

    result = graph.request("POST", "/me/contacts", account_id, json=contact)
    if not result:
        raise ValueError("Failed to create contact")
    return result


@mcp.tool
def update_contact(
    contact_id: str, updates: dict[str, Any], account_id: str
) -> dict[str, Any]:
    """Update contact information"""
    result = graph.request(
        "PATCH", f"/me/contacts/{contact_id}", account_id, json=updates
    )
    return result or {"status": "updated"}


@mcp.tool
def delete_contact(contact_id: str, account_id: str) -> dict[str, str]:
    """Delete a contact"""
    graph.request("DELETE", f"/me/contacts/{contact_id}", account_id)
    return {"status": "deleted"}


@mcp.tool
def list_files(
    account_id: str, path: str = "/", limit: int = 50
) -> list[dict[str, Any]]:
    """List files and folders in OneDrive"""
    endpoint = (
        "/me/drive/root/children"
        if path == "/"
        else f"/me/drive/root:/{path}:/children"
    )
    params = {
        "$top": min(limit, 100),
        "$select": "id,name,size,lastModifiedDateTime,folder,file,@microsoft.graph.downloadUrl",
    }

    items = list(
        graph.request_paginated(endpoint, account_id, params=params, limit=limit)
    )

    return [
        {
            "id": item["id"],
            "name": item["name"],
            "type": "folder" if "folder" in item else "file",
            "size": item.get("size", 0),
            "modified": item.get("lastModifiedDateTime"),
            "download_url": item.get("@microsoft.graph.downloadUrl"),
        }
        for item in items
    ]


@mcp.tool
def get_file(file_id: str, account_id: str) -> dict[str, Any]:
    """Get file content from OneDrive and return it in the result
    
    Args:
        file_id: The OneDrive file ID
        account_id: The account ID to use
    
    Returns:
        Dictionary with file metadata and content:
        - For text files: content is returned as a string
        - For binary files: content is returned as base64-encoded string
    """
    metadata = graph.request("GET", f"/me/drive/items/{file_id}", account_id)
    if not metadata:
        raise ValueError(f"File with ID {file_id} not found")

    # Get file content using the /content endpoint
    try:
        content_bytes = graph.download_raw(
            f"/me/drive/items/{file_id}/content",
            account_id
        )
    except Exception as e:
        raise ValueError(f"Failed to download file content: {e}")

    # Get mime type and file name to determine how to process
    mime_type = metadata.get("file", {}).get("mimeType", "application/octet-stream")
    file_name = metadata.get("name", "").lower()
    
    # Check if it's a Word document (.docx)
    is_word_doc = (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or file_name.endswith(".docx")
    )
    
    # Extract text from Word documents
    if is_word_doc:
        try:
            # Load Word document from bytes
            doc = Document(io.BytesIO(content_bytes))
            # Extract all text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs]
            content = "\n".join(paragraphs)
            is_text = True
            encoding = "utf-8"
        except Exception:
            # If extraction fails, fall back to base64
            content = base64.b64encode(content_bytes).decode("utf-8")
            is_text = False
            encoding = "base64"
    else:
        # Determine if it's a text file based on mime type
        text_mime_types = [
            "text/",
            "application/json",
            "application/xml",
            "application/javascript",
            "application/x-sh",
            "application/x-python",
        ]
        is_text = any(mime_type.startswith(prefix) for prefix in text_mime_types)
        
        # Return content as text or base64
        if is_text:
            try:
                # Try to decode as UTF-8
                content = content_bytes.decode("utf-8")
                encoding = "utf-8"
            except UnicodeDecodeError:
                # If UTF-8 fails, encode as base64
                content = base64.b64encode(content_bytes).decode("utf-8")
                is_text = False
                encoding = "base64"
        else:
            # Binary file - encode as base64
            content = base64.b64encode(content_bytes).decode("utf-8")
            encoding = "base64"

    return {
        "id": file_id,
        "name": metadata.get("name", "unknown"),
        "size": metadata.get("size", 0),
        "size_mb": round(metadata.get("size", 0) / (1024 * 1024), 2),
        "mime_type": mime_type,
        "content": content,
        "is_text": is_text,
        "encoding": encoding,
    }


@mcp.tool
def create_file(
    onedrive_path: str, file_name: str, file_content: str, account_id: str
) -> dict[str, Any]:
    """Create a Word document in OneDrive from text content
    
    Args:
        onedrive_path: The OneDrive folder path (e.g., "/Documents" or "/")
        file_name: Name of the file (without .docx extension, will be added automatically)
        file_content: The text content to include in the Word document
        account_id: The account ID to use
    """
    # Ensure file_name has .docx extension
    if not file_name.endswith('.docx'):
        file_name = f"{file_name}.docx"
    
    # Create Word document in memory
    doc = Document()
    # Add content to document (split by newlines to create paragraphs)
    for line in file_content.split('\n'):
        doc.add_paragraph(line)
    
    # Save document to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    data = doc_bytes.read()
    
    # Construct full path
    if onedrive_path == "/" or onedrive_path == "":
        full_path = f"/{file_name}"
    else:
        full_path = f"{onedrive_path.rstrip('/')}/{file_name}"
    
    result = graph.upload_large_file(
        f"/me/drive/root:{full_path}:", data, account_id
    )
    if not result:
        raise ValueError(f"Failed to create file at path: {full_path}")
    return result


@mcp.tool
def update_file(
    file_id: str, file_name: str, file_content: str, account_id: str
) -> dict[str, Any]:
    """Update a OneDrive file with new Word document content and optionally rename it
    
    Args:
        file_id: The OneDrive file ID to update
        file_name (Optional): Name of the file (without .docx extension, will be added automatically)
        file_content (Optional): The text content to include in the Word document
        account_id: The account ID to use
    """
    # Ensure file_name has .docx extension
    if not file_name.endswith('.docx'):
        file_name = f"{file_name}.docx"
    
    # Get current file metadata to check if name needs updating
    current_file = graph.request("GET", f"/me/drive/items/{file_id}", account_id)
    if not current_file:
        raise ValueError(f"File with ID {file_id} not found")
    
    current_name = current_file.get("name", "")
    name_changed = current_name != file_name
    
    # Create Word document in memory
    doc = Document()
    # Add content to document (split by newlines to create paragraphs)
    for line in file_content.split('\n'):
        doc.add_paragraph(line)
    
    # Save document to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    data = doc_bytes.read()
    
    # Update file content using PUT to /content endpoint
    # For small files, use direct PUT. For large files, use upload session.
    # The chunk size is 4.5MB (15 * 320 KiB)
    file_size = len(data)
    UPLOAD_CHUNK_SIZE = 15 * 320 * 1024  # 4,915,200 bytes
    
    if file_size <= UPLOAD_CHUNK_SIZE:
        # Direct PUT for small files - this is the most reliable method for updates
        # Using explicit PUT to /content endpoint ensures the file content is replaced
        result = graph.request(
            "PUT",
            f"/me/drive/items/{file_id}/content",
            account_id,
            data=data
        )
    else:
        # For large files, use upload session approach
        result = graph.upload_large_file(
            f"/me/drive/items/{file_id}",
            data,
            account_id
        )
    
    if not result:
        raise ValueError(f"Failed to update file with ID: {file_id}")
    
    # Update file name if it changed
    if name_changed:
        name_update = graph.request(
            "PATCH",
            f"/me/drive/items/{file_id}",
            account_id,
            json={"name": file_name}
        )
        if name_update:
            result["name"] = name_update.get("name", file_name)
    
    return result


@mcp.tool
def delete_file(file_id: str, account_id: str) -> dict[str, str]:
    """Delete a file or folder"""
    graph.request("DELETE", f"/me/drive/items/{file_id}", account_id)
    return {"status": "deleted"}


@mcp.tool
def get_attachment(
    email_id: str, attachment_id: str, save_path: str, account_id: str
) -> dict[str, Any]:
    """Download email attachment to a specified file path"""
    result = graph.request(
        "GET", f"/me/messages/{email_id}/attachments/{attachment_id}", account_id
    )

    if not result:
        raise ValueError("Attachment not found")

    if "contentBytes" not in result:
        raise ValueError("Attachment content not available")

    # Save attachment to file
    path = pl.Path(save_path).expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)
    content_bytes = base64.b64decode(result["contentBytes"])
    path.write_bytes(content_bytes)

    return {
        "name": result.get("name", "unknown"),
        "content_type": result.get("contentType", "application/octet-stream"),
        "size": result.get("size", 0),
        "saved_to": str(path),
    }


@mcp.tool
def search_files(
    query: str,
    account_id: str,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Search for files in OneDrive using the modern search API."""
    items = list(graph.search_query(query, ["driveItem"], account_id, limit))

    return [
        {
            "id": item["id"],
            "name": item["name"],
            "type": "folder" if "folder" in item else "file",
            "size": item.get("size", 0),
            "modified": item.get("lastModifiedDateTime"),
            "download_url": item.get("@microsoft.graph.downloadUrl"),
        }
        for item in items
    ]


@mcp.tool
def search_emails(
    query: str,
    account_id: str,
    limit: int = 50,
    folder: str | None = None,
) -> list[dict[str, Any]]:
    """Search emails using the modern search API."""
    if folder:
        # For folder-specific search, use the traditional endpoint
        folder_path = FOLDERS.get(folder.casefold(), folder)
        endpoint = f"/me/mailFolders/{folder_path}/messages"

        params = {
            "$search": f'"{query}"',
            "$top": min(limit, 100),
            "$select": "id,subject,from,toRecipients,receivedDateTime,hasAttachments,body,conversationId,isRead",
        }

        return list(
            graph.request_paginated(endpoint, account_id, params=params, limit=limit)
        )

    return list(graph.search_query(query, ["message"], account_id, limit))


@mcp.tool
def search_events(
    query: str,
    account_id: str,
    days_ahead: int = 365,
    days_back: int = 365,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Search calendar events using the modern search API."""
    events = list(graph.search_query(query, ["event"], account_id, limit))

    # Filter by date range if needed
    if days_ahead != 365 or days_back != 365:
        now = dt.datetime.now(dt.timezone.utc)
        start = now - dt.timedelta(days=days_back)
        end = now + dt.timedelta(days=days_ahead)

        filtered_events = []
        for event in events:
            event_start = dt.datetime.fromisoformat(
                event.get("start", {}).get("dateTime", "").replace("Z", "+00:00")
            )
            event_end = dt.datetime.fromisoformat(
                event.get("end", {}).get("dateTime", "").replace("Z", "+00:00")
            )

            if event_start <= end and event_end >= start:
                filtered_events.append(event)

        return filtered_events

    return events


@mcp.tool
def search_contacts(
    query: str,
    account_id: str,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Search contacts. Uses traditional search since unified_search doesn't support contacts."""
    params = {
        "$search": f'"{query}"',
        "$top": min(limit, 100),
    }

    contacts = list(
        graph.request_paginated("/me/contacts", account_id, params=params, limit=limit)
    )

    return contacts


@mcp.tool
def unified_search(
    query: str,
    account_id: str,
    entity_types: list[str] | None = None,
    limit: int = 50,
) -> dict[str, list[dict[str, Any]]]:
    """Search across multiple Microsoft 365 resources using the modern search API

    entity_types can include: 'message', 'event', 'drive', 'driveItem', 'list', 'listItem', 'site'
    If not specified, searches across all available types.
    """
    if not entity_types:
        entity_types = ["message", "event", "driveItem"]

    results = {entity_type: [] for entity_type in entity_types}

    items = list(graph.search_query(query, entity_types, account_id, limit))

    for item in items:
        resource_type = item.get("@odata.type", "").split(".")[-1]

        if resource_type == "message":
            results.setdefault("message", []).append(item)
        elif resource_type == "event":
            results.setdefault("event", []).append(item)
        elif resource_type in ["driveItem", "file", "folder"]:
            results.setdefault("driveItem", []).append(item)
        else:
            results.setdefault("other", []).append(item)

    return {k: v for k, v in results.items() if v}
